# build a CV creator application, including a section in it where you can showcase the list of skills with bullet points
# hint: print("\u2022") prints out bullet points or p.style = "List Bullet" performs the same operation

# if the line below gives an error, then docx package is not installed on your computer. Go to terminal and type "pip3 install python-docx"
from docx import Document
# the line below is added to resize pictures
from docx.shared import Inches

# invoke the related Document functions
document = Document()

# make sure the picture you want to add to your CV is in the same folder as the application file, then write the name of the picture in line 14 between " "
document.add_picture(
    "PP.jpg", width=Inches(4))

# details to add to your CV
name = input("Enter your name : ")
phone_num = input("Enter your phone number : ")
e_mail = input("Enter your e-mail : ")

document.add_paragraph(
    name + " | " + phone_num + " | " + e_mail)

# about me
document.add_heading("About Me")
#about_me = input("Tell me about yourself: ")
document.add_paragraph(input("Tell about yourself: "))

# work experience
document.add_heading("Work Experince")
p = document.add_paragraph()

company = input("Enter company: ")
from_date = input("From Date ")
to_date = input("To Date ")

p.add_run(company + " " ).bold = True
p.add_run(from_date + "-" + to_date + "\n").italic = True
experience_details =input(
    "Describe your experience at " + company + " ")
p.add_run(experience_details)

# in case you want to add more experiences
while True:
    has_more_experiences = input(
        "Do you have more experiences? (Yes or No): ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input("From Date ")
        to_date = input("To Date ")

        p.add_run(company + " " ).bold = True
        p.add_run(from_date + "-" + to_date + "\n").italic = True
        experience_details =input(
            "Describe your experience at " + company + " ")
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading("Skills")
p = document.add_paragraph()
skill = input("Enter skill: ")
level = input("Enter your level a scale from 0 to 5: ")
certificate = input("Any certificate? (Yes or No): ")

p.style ="List Bullet"

p.add_run(skill + ", Level: " + level + ", Certificate: " + certificate ).bold = True

# in case you want to add more skills
while True:
    has_more_skills = input(
        "Do you have more skills? (Yes or No): ")
    if has_more_skills.lower() == "yes":
        p = document.add_paragraph()

        skill = input("Enter skill: ")
        level = input("Enter your level a scale from 0 to 5: ")
        certificate = input("Any certificate? (Yes or No): ")
        
        p.add_run(skill + ", Level: " + level + ", Certificate: " + certificate ).bold = True
        p.style ="List Bullet"
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Şerife's first Python application."

# saving on the document, make sure the document is in the same folder as well
document.save("CV.docx")